# """
# Resume parsing functionality for PDF and DOCX files
# """

# import re
# import json
# from datetime import datetime
# from pathlib import Path
# import PyPDF2
# from docx import Document
# import nltk
# from nltk.tokenize import word_tokenize, sent_tokenize
# from config import SKILLS_KEYWORDS, EDUCATION_KEYWORDS

# # Download required NLTK data
# try:
#     nltk.data.find('tokenizers/punkt')
# except LookupError:
#     nltk.download('punkt', quiet=True)

# class ResumeParser:
#     def __init__(self):
#         self.skills_keywords = [skill.lower() for skill in SKILLS_KEYWORDS]
#         self.education_keywords = [edu.lower() for edu in EDUCATION_KEYWORDS]
        
#     def extract_text_from_pdf(self, file_path):
#         """Extract text from PDF file"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text
#         except Exception as e:
#             print(f"Error extracting text from PDF: {e}")
#             return ""
    
#     def extract_text_from_docx(self, file_path):
#         """Extract text from DOCX file"""
#         try:
#             doc = Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text
#         except Exception as e:
#             print(f"Error extracting text from DOCX: {e}")
#             return ""
    
#     def extract_name(self, text):
#         """Extract name from resume text"""
#         lines = text.split('\n')[:5]  # Check first 5 lines
        
#         # Common patterns for names
#         name_patterns = [
#             r'^([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
#             r'Name[:\s]*([A-Z][a-z]+\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)',
#             r'^([A-Z]{2,}\s+[A-Z]{2,}(?:\s+[A-Z]{2,})?)'
#         ]
        
#         for line in lines:
#             line = line.strip()
#             if len(line) > 0:
#                 for pattern in name_patterns:
#                     match = re.search(pattern, line)
#                     if match:
#                         return match.group(1).strip()
        
#         return "Name not found"
    
#     def extract_email(self, text):
#         """Extract email from resume text"""
#         email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
#         emails = re.findall(email_pattern, text)
#         return emails[0] if emails else "Email not found"
    
#     def extract_phone(self, text):
#         """Extract phone number from resume text"""
#         phone_patterns = [
#             r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
#             r'\+?\d{10,15}',
#             r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'
#         ]
        
#         for pattern in phone_patterns:
#             phones = re.findall(pattern, text)
#             if phones:
#                 return phones[0]
        
#         return "Phone not found"
    
#     def extract_skills(self, text):
#         """Extract skills from resume text"""
#         text_lower = text.lower()
#         found_skills = []
        
#         for skill in self.skills_keywords:
#             if skill in text_lower:
#                 found_skills.append(skill.title())
        
#         # Also look for skills in a dedicated skills section
#         skills_section_pattern = r'skills?[:\s]*([^\n]+(?:\n[^\n]+)*?)(?:\n\s*\n|\Z)'
#         skills_match = re.search(skills_section_pattern, text, re.IGNORECASE | re.MULTILINE)
        
#         if skills_match:
#             skills_text = skills_match.group(1)
#             additional_skills = re.findall(r'\b[A-Za-z][A-Za-z\s.+#-]*\b', skills_text)
#             found_skills.extend([skill.strip().title() for skill in additional_skills if len(skill.strip()) > 2])
        
#         return list(set(found_skills))
    
#     def extract_experience(self, text):
#         """Extract work experience from resume text"""
#         experience_patterns = [
#             r'experience[:\s]*(\d+)\s*(?:years?|yrs?)',
#             r'(\d+)\s*(?:years?|yrs?)\s*(?:of\s*)?experience',
#             r'work\s*experience[:\s]*(\d+)\s*(?:years?|yrs?)'
#         ]
        
#         for pattern in experience_patterns:
#             match = re.search(pattern, text, re.IGNORECASE)
#             if match:
#                 return int(match.group(1))
        
#         # Try to count job positions
#         job_patterns = [
#             r'(?:software\s+engineer|developer|analyst|manager|consultant)',
#             r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{4}\s*[-–]\s*(?:present|current|\d{4})',
#             r'\d{4}\s*[-–]\s*(?:present|current|\d{4})'
#         ]
        
#         job_count = 0
#         for pattern in job_patterns:
#             matches = re.findall(pattern, text, re.IGNORECASE)
#             job_count += len(matches)
        
#         return max(1, job_count // 2) if job_count > 0 else 0
    
#     def extract_education(self, text):
#         """Extract education information from resume text"""
#         education = []
        
#         # Look for degree patterns
#         degree_patterns = [
#             r'(bachelor[^,\n]*|master[^,\n]*|phd[^,\n]*|doctorate[^,\n]*)',
#             r'(b\.?tech[^,\n]*|m\.?tech[^,\n]*|mba[^,\n]*|bca[^,\n]*|mca[^,\n]*)',
#             r'(be[^,\n]*|me[^,\n]*|ms[^,\n]*|bs[^,\n]*)'
#         ]
        
#         for pattern in degree_patterns:
#             matches = re.findall(pattern, text, re.IGNORECASE)
#             education.extend([match.strip().title() for match in matches])
        
#         return list(set(education)) if education else ["Education details not found"]
    
#     def parse_resume(self, file_path):
#         """Parse a single resume file"""
#         file_path = Path(file_path)
        
#         # Extract text based on file type
#         if file_path.suffix.lower() == '.pdf':
#             text = self.extract_text_from_pdf(file_path)
#         elif file_path.suffix.lower() == '.docx':
#             text = self.extract_text_from_docx(file_path)
#         else:
#             return None
        
#         if not text:
#             return None
        
#         # Extract structured information
#         parsed_data = {
#             "file_name": file_path.name,
#             "parsed_date": datetime.now().isoformat(),
#             "name": self.extract_name(text),
#             "email": self.extract_email(text),
#             "phone": self.extract_phone(text),
#             "skills": self.extract_skills(text),
#             "experience_years": self.extract_experience(text),
#             "education": self.extract_education(text),
#             "raw_text": text[:2000]  # Store first 2000 chars for reference
#         }
        
#         return parsed_data
    
#     def save_parsed_resume(self, parsed_data, output_dir):
#         """Save parsed resume data to JSON file"""
#         output_dir = Path(output_dir)
#         output_dir.mkdir(exist_ok=True)
        
#         # Create filename based on extracted name or original filename
#         name_for_file = parsed_data.get('name', 'unknown').replace(' ', '_').replace('/', '_')
#         json_filename = f"{name_for_file}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
#         json_path = output_dir / json_filename
        
#         try:
#             with open(json_path, 'w', encoding='utf-8') as f:
#                 json.dump(parsed_data, f, indent=2, ensure_ascii=False)
#             return json_path
#         except Exception as e:
#             print(f"Error saving parsed resume: {e}")
#             return import re

import re
import json
from datetime import datetime
from pathlib import Path
import PyPDF2
from docx import Document
import google.generativeai as genai
from config import DEFAULT_PARSING_PROMPT

class ResumeParser:
    def __init__(self, api_key):
        """Initialize Gemini AI client with the free-tier model"""
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel("gemini-1.5-flash")  # Use the free-tier model

    def extract_text_from_pdf(self, file_path):
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def extract_text_from_docx(self, file_path):
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def extract_text_from_txt(self, file_path):
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Error extracting text from TXT: {e}")
            return ""
    
    def extract_text(self, file_path):
        """Extract text from file based on extension"""
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            return ""
    
    def parse_with_gemini(self, text, filename):
        """Parse resume text using Gemini AI"""
        try:
            # Enhanced prompt with filename context
            enhanced_prompt = f"""
{DEFAULT_PARSING_PROMPT}

RESUME TEXT FROM FILE '{filename}':
{text[:15000]}  # Limit text length
"""
            
            response = self.model.generate_content(enhanced_prompt)
            
            # Try to extract JSON from response
            response_text = response.text.strip()
            
            # Clean the response text
            response_text = self.clean_json_response(response_text)
            
            # Parse JSON
            parsed_data = json.loads(response_text)
            
            # Validate and ensure all required fields are present
            parsed_data = self.validate_parsed_data(parsed_data, text)
            
            return parsed_data
        except Exception as e:
            print(f"Error parsing with Gemini: {e}")
            return None
    
    def clean_json_response(self, response_text):
        """Clean the JSON response from Gemini"""
        # Remove markdown code blocks if present
        if response_text.startswith('```json'):
            response_text = response_text[7:]
        if response_text.startswith('```'):
            response_text = response_text[3:]
        if response_text.endswith('```'):
            response_text = response_text[:-3]
        
        # Remove any text before the first {
        start_idx = response_text.find('{')
        if start_idx > 0:
            response_text = response_text[start_idx:]
        
        # Remove any text after the last }
        end_idx = response_text.rfind('}')
        if end_idx < len(response_text) - 1:
            response_text = response_text[:end_idx + 1]
        
        return response_text.strip()
    
    def validate_parsed_data(self, parsed_data, text):
        """Validate and ensure all required fields are present in parsed data"""
        # Ensure all required fields are present
        required_fields = ['name', 'email', 'phone', 'experience_years', 'skills', 'education', 'summary']
        
        for field in required_fields:
            if field not in parsed_data:
                if field == 'skills' or field == 'education':
                    parsed_data[field] = []
                elif field == 'experience_years':
                    parsed_data[field] = 0
                else:
                    parsed_data[field] = ""
        
        # Validate name field specifically
        if not parsed_data['name'] or parsed_data['name'].lower() in ['unknown', 'not found', 'n/a']:
            parsed_data['name'] = self.extract_name_fallback(text)
        
        # Validate email field
        if not parsed_data['email'] or parsed_data['email'].lower() in ['unknown', 'not found', 'n/a']:
            parsed_data['email'] = self.extract_email_fallback(text)
        
        # Validate phone field
        if not parsed_data['phone'] or parsed_data['phone'].lower() in ['unknown', 'not found', 'n/a']:
            parsed_data['phone'] = self.extract_phone_fallback(text)
        
        return parsed_data
    
    def extract_name_fallback(self, text):
        """Fallback method to extract name from text"""
        # Look for common name patterns
        lines = text.split('\n')
        for line in lines[:10]:  # Check first 10 lines
            line = line.strip()
            # Look for lines that might contain names (2-3 words, title case)
            if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}$', line):
                return line
            # Look for lines with name prefixes
            if any(prefix in line.lower() for prefix in ['name:', 'full name:', 'candidate:']):
                parts = re.split(r'name:|full name:|candidate:', line, flags=re.IGNORECASE)
                if len(parts) > 1:
                    name = parts[1].strip()
                    if name:
                        return name
        
        return "Unknown Candidate"
    
    def extract_email_fallback(self, text):
        """Fallback method to extract email from text"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        return emails[0] if emails else "Not found"
    
    def extract_phone_fallback(self, text):
        """Fallback method to extract phone from text"""
        phone_patterns = [
            r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            r'\+?\d{10,15}',
            r'\(\d{3}\)\s?\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                return phones[0]
        
        return "Not found"
    
    def parse_resume(self, file_path):
        """Parse a single resume file"""
        file_path = Path(file_path)
        
        # Extract text from file
        text = self.extract_text(file_path)
        
        if not text:
            print(f"Failed to extract text from {file_path.name}")
            return None
        
        # Parse with Gemini AI
        parsed_data = self.parse_with_gemini(text, file_path.name)
        
        if parsed_data is None:
            print(f"Failed to parse resume with Gemini AI: {file_path.name}")
            return None
        
        # Add metadata
        parsed_data.update({
            "file_name": file_path.name,
            "parsed_date": datetime.now().isoformat(),
            "raw_text_preview": text[:500]  # Store first 500 chars for reference
        })
        
        return parsed_data